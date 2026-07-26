import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import io

# 1. CONFIGURACIÓN DE LA PÁGINA
st.set_page_config(page_title="Sistema EVM - Consultoría", layout="wide")
st.title("📊 Sistema de Análisis de Valor Ganado (EVM)")
st.write("Herramienta profesional para la optimización y control de proyectos en ingeniería.")

# 2. INGRESO DINÁMICO DE DATOS DEL PROYECTO (Barra Lateral)
st.sidebar.header("Configuración del Proyecto")
nombre_proyecto = st.sidebar.text_input("Nombre del proyecto:", "Proyecto de Infraestructura")
BAC = st.sidebar.number_input("Presupuesto a la Conclusión (BAC) en $:", min_value=0.0, value=10000.0, step=1000.0)
duracion_total = st.sidebar.number_input("Duración total planificada (periodos):", min_value=1, value=12)
periodos_transcurridos = st.sidebar.number_input("Periodos transcurridos a la fecha:", min_value=1, max_value=int(duracion_total), value=1)

st.subheader("📝 Ingreso de Datos Acumulados por Periodo")
st.write("Ingrese los porcentajes de avance (0 a 100) y los costos reales incurridos. Edite la tabla directamente:")

# Crear tabla dinámica para ingreso de datos
datos_iniciales = {
    "Periodo": [i for i in range(1, int(periodos_transcurridos) + 1)],
    "Avance_Planeado_Acumulado_%": [0.0] * int(periodos_transcurridos),
    "Avance_Real_Acumulado_%": [0.0] * int(periodos_transcurridos),
    "Costo_Real_Acumulado_$": [0.0] * int(periodos_transcurridos)
}
df_input = pd.DataFrame(datos_iniciales)
df_editado = st.data_editor(df_input, hide_index=True, use_container_width=True)

if st.button("Ejecutar Análisis EVM"):
    # Procesamiento de la matriz
    avance_planeado = df_editado["Avance_Planeado_Acumulado_%"] / 100.0
    avance_real = df_editado["Avance_Real_Acumulado_%"] / 100.0
    costo_real = df_editado["Costo_Real_Acumulado_$"]
    
    df = pd.DataFrame({
        "Periodo": df_editado["Periodo"],
        "Avance_Planeado_Acumulado": avance_planeado,
        "Avance_Real_Acumulado": avance_real,
        "Costo_Real_Acumulado": costo_real
    })
    
    # 3. CÁLCULO DE FÓRMULAS A LA FECHA DE CORTE
    avance_planeado_actual = df["Avance_Planeado_Acumulado"].iloc[-1]
    avance_real_actual = df["Avance_Real_Acumulado"].iloc[-1]
    
    PV = BAC * avance_planeado_actual
    AC = df["Costo_Real_Acumulado"].iloc[-1]
    EV = BAC * avance_real_actual
    
    # Índices de Rendimiento
    CPI = EV / AC if AC > 0 else 0
    SPI = EV / PV if PV > 0 else 0
    
    # Variaciones
    CV = EV - AC
    SV = EV - PV
    
    # Proyecciones
    EAC = BAC / CPI if CPI > 0 else BAC
    ETC = (BAC - EV) / CPI if CPI > 0 else BAC

    # 4. INTERPRETACIÓN DINÁMICA DE INDICADORES
    if CPI > 1:
        int_cpi = "Eficiente. Los costos planificados son mayores a los costos reales."
    elif CPI == 1:
        int_cpi = "Alineado. Los costos planificados son iguales a los costos reales."
    else:
        int_cpi = "Ineficiencia. Los costos planificados son menores a los costos reales."
        
    if SPI > 1:
        int_spi = "Acelerado. Se avanza a un ritmo mayor al planificado."
    elif SPI == 1:
        int_spi = "A tiempo. Se avanza a un ritmo igual al planificado."
    else:
        int_spi = "Ineficiencia. Se avanza a un ritmo menor al planificado."
        
    if CV > 0:
        int_cv = "Favorable. Costos planificados mayores a reales."
    elif CV == 0:
        int_cv = "Neutro. Costos planificados iguales a reales."
    else:
        int_cv = "Ineficiencia. Costos planificados menores a reales."
        
    if SV > 0:
        int_sv = "Favorable. El proyecto está adelantado respecto al cronograma base."
    elif SV == 0:
        int_sv = "Neutro. El proyecto está exactamente a tiempo."
    else:
        int_sv = "Desfavorable. El proyecto está atrasado respecto al cronograma base."

    # Mostrar Resultados en Pantalla
    st.markdown("---")
    st.header("📈 Resultados del Análisis")
    
    col1, col2, col3 = st.columns(3)
    col1.metric("Valor Planificado (PV)", f"${PV:,.2f}")
    col2.metric("Costo Real (AC)", f"${AC:,.2f}")
    col3.metric("Valor Ganado (EV)", f"${EV:,.2f}")
    
    col4, col5 = st.columns(2)
    col4.metric("CPI (Costo)", f"{CPI:.2f}", int_cpi)
    col5.metric("SPI (Cronograma)", f"{SPI:.2f}", int_spi)

    # 5. GRÁFICA DE VALOR GANADO (Curvas "S")
    df['PV_Acum'] = BAC * df['Avance_Planeado_Acumulado']
    df['EV_Acum'] = BAC * df['Avance_Real_Acumulado']
    df['AC_Acum'] = df['Costo_Real_Acumulado']
    
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.plot(df['Periodo'], df['PV_Acum'], marker='o', label='Valor Planificado (PV)', color='blue')
    ax.plot(df['Periodo'], df['EV_Acum'], marker='o', label='Valor Ganado (EV)', color='green')
    ax.plot(df['Periodo'], df['AC_Acum'], marker='o', label='Costo Real (AC)', color='red')
    
    ax.set_xlim(1, duracion_total)
    ax.set_title(f'Curvas "S" de Coste - {nombre_proyecto}')
    ax.set_xlabel('Tiempo (Periodos)')
    ax.set_ylabel('Costo Acumulado ($)')
    ax.grid(True, linestyle='--', alpha=0.7)
    ax.legend()
    
    st.pyplot(fig)

    # 6. GENERACIÓN DEL INFORME INTEGRAL EN WORD
    doc = Document()
    doc.add_heading(f'Informe Integral EVM: {nombre_proyecto}', 0)
    doc.add_paragraph('Variables y Limitaciones:')
    doc.add_paragraph(f'- Presupuesto a la conclusión (BAC): ${BAC:,.2f}')
    doc.add_paragraph(f'- Duración planificada: {duracion_total} periodos')
    doc.add_paragraph(f'- Fecha de corte analizada: Periodo {periodos_transcurridos}')
    
    doc.add_heading('Resultados e Interpretación de Métricas', level=1)
    doc.add_paragraph(f'Valor Planificado (PV): ${PV:,.2f}')
    doc.add_paragraph(f'Costo Real (AC): ${AC:,.2f}')
    doc.add_paragraph(f'Valor Ganado (EV): ${EV:,.2f}')
    
    doc.add_heading('Índices de Rendimiento', level=2)
    doc.add_paragraph(f'CPI: {CPI:.2f} -> {int_cpi}')
    doc.add_paragraph(f'SPI: {SPI:.2f} -> {int_spi}')
    
    doc.add_heading('Variaciones del Proyecto', level=2)
    doc.add_paragraph(f'Variación de Costo (CV): ${CV:,.2f} -> {int_cv}')
    doc.add_paragraph(f'Variación de Cronograma (SV): ${SV:,.2f} -> {int_sv}')
    
    doc.add_heading('Proyecciones', level=2)
    doc.add_paragraph(f'Estimación a la Conclusión (EAC): ${EAC:,.2f}')
    doc.add_paragraph(f'Estimación hasta la Conclusión (ETC): ${ETC:,.2f}')
    
    # Guardar gráfico temporalmente para insertarlo en Word
    img_stream = io.BytesIO()
    fig.savefig(img_stream, format='png', bbox_inches='tight')
    img_stream.seek(0)
    doc.add_heading('Gráfica de Valor Ganado (Curvas S)', level=1)
    doc.add_picture(img_stream, width=Inches(6.0))
    
    doc.add_heading('Conclusión Ejecutiva', level=1)
    # Lógica condicional para detectar el falso ahorro por sub-ejecución
    if CPI > 1 and SPI < 1:
        conclusion_ejecutiva = (
            "ESTATUS INTEGRAL DEL PROYECTO: ALERTA DE SUB-EJECUCIÓN\n\n"
            f"Existe una falsa apariencia de ahorro. El proyecto refleja una 'eficiencia' en costos (CPI: {CPI:.2f}), "
            f"pero es una consecuencia directa del atraso crítico en el cronograma (SPI: {SPI:.2f}). "
            f"El bajo gasto real (AC: ${AC:,.2f}) se debe a que no se ha ejecutado el volumen de obra programado "
            f"para la fecha (PV: ${PV:,.2f}).\n\n"
            f"Impacto y Acción Requerida: Para recuperar la variación de cronograma de ${SV:,.2f}, "
            "la gerencia deberá inyectar recursos (horas extras, acelerar suministros de materiales o integrar subcontratistas adicionales). "
            "Estas medidas de 'crashing' (compresión) encarecerán la mano de obra y diluirán rápidamente el actual margen positivo de costos. "
            f"Por lo tanto, la estimación a la conclusión (EAC: ${EAC:,.2f}) es engañosa y no se sostendrá una vez que se acelere el ritmo físico de la obra."
        )
    else:
        conclusion_ejecutiva = f"ESTATUS ACTUAL DEL PROYECTO:\n- Costo: {int_cpi}\n- Tiempo: {int_sv}"
        
    doc.add_paragraph(conclusion_ejecutiva)
    
    # Preparar el archivo Word para descarga directa en la App
    word_stream = io.BytesIO()
    doc.save(word_stream)
    word_stream.seek(0)
    
    st.markdown("---")
    st.header("📄 Informe Listo para Descargar")
    st.download_button(
        label="Descargar Informe Ejecutivo en Word",
        data=word_stream,
        file_name=f"Informe_EVM_{nombre_proyecto.replace(' ', '_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )